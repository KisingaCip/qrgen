"""
Generate qr codes with superimposed images
Dependencies: PIL or pillow (I use pillow), qrcode
"""

import qrcode
import sys
import argparse
from PIL import Image, ImageEnhance
from PIL import ImageFont
from PIL import ImageDraw
import glob, shutil, os
from docx import Document
from docx.shared import Cm
import time
from docx.enum.text import WD_ALIGN_PARAGRAPH

# prepare the the code range and a temp folder

try:
    start= raw_input("First Barcode:_ ")
    stop = raw_input('Last Barcode:_ ')
except:
    start= input("First Barcode:_ ") 
    stop = input('Last Barcode:_ ') 
start_time =time.time()
try:
    if int(stop)>=int(start):
        subdir= 'temp'
        if os.path.isdir(subdir):
            shutil.rmtree(subdir)
        os.mkdir(subdir)
        os.chdir(subdir)
        print("Generating codes......  please wait!")
    else:
        print("Invalid range! Exiting....!")
except:
    sys.exit()

def do_qr(text, mark=None, box_size=10, border=4, fileext="png"):
    """ Writes out a QR code to a PNG file """
    qr = qrcode.QRCode(
        error_correction=qrcode.constants.ERROR_CORRECT_H,  # high error correction (allows image superimpose)
        box_size=box_size,  # pixels per "box" of the QR
        border=border,  # min 4. border size in pixels
    )
   
    qr.add_data(text)
    qr.make(fit=True)
   
    img = qr.make_image()

    if mark is not None:
        overpos = img.size[0] / 2 - mark.size[0] / 2
        img = watermark(img, mark, (overpos, overpos))

    return img

# from http://code.activestate.com/recipes/362879-watermark-with-pil/
def watermark(im, mark, position):
    """Adds a watermark to an image."""
    if im.mode != 'RGBA':
        im = im.convert('RGBA')
    # create a transparent layer the size of the image and draw the
    # watermark in that layer.
    layer = Image.new('RGBA', im.size, (0, 0, 0, 0))
    if position == 'tile':
        for y in range(0, im.size[1], mark.size[1]):
            for x in range(0, im.size[0], mark.size[0]):
                layer.paste(mark, (x, y))
    elif position == 'scale':
        # scale, but preserve the aspect ratio
        ratio = min(
            float(im.size[0]) / mark.size[0], float(im.size[1]) / mark.size[1])
        w = int(mark.size[0] * ratio)
        h = int(mark.size[1] * ratio)
        mark = mark.resize((w, h))
        layer.paste(mark, ((im.size[0] - w) / 2, (im.size[1] - h) / 2))
    else:
        layer.paste(mark, position)
    # composite the watermark with the layer
    return Image.composite(layer, im, layer)


if __name__ == "__main__":
    # Set up command line args
    parser = argparse.ArgumentParser(fromfile_prefix_chars='@',
                                        description='Generate QR codes as PNG files',
                                        epilog='This tool is provided AS-IS and no support or warranty is given or implied. ')
    parser.add_argument('text', nargs='*', default="", help='Text to encode in the QR code.')
    parser.add_argument('--overlay', default=None, help='a PNG file to overlay in the centre of the code. Test your codes work!')
    parser.add_argument('--prefix', '--url', '-p', default="", help='Optional text added to the start of the code\'s text, e.g. http://server/qr/')
    parser.add_argument('--suffix', '-s', default="", help='Optional text added to the end of the code\'s text.')
    parser.add_argument('--format', '-f', default="png", choices=["png", "jpg", "bmp"], help='Image format (default: PNG)')
    parser.add_argument('--size', default=5, type=int, help='Pixel width per "box" of the QR (default: 5)')
    parser.add_argument('--border', default=4, type=int, help='Border width in "boxes" (at least 4 is required by the standards)')
    args = parser.parse_args()

    # Open the watermark file
    mark = None
    if args.overlay not in (None, ''):
        try:
            print "Opening overlay file {0}".format(args.overlay)
            mark = Image.open(args.overlay)
        except Exception as x:
            print "Note: Couldn't open overlay {0}, is it a PNG?" % args.overlay, x

    for line in range(int(start),int(stop)+1):
            # Accept stdin lines
        line ='prefix%06d' %(line)

        #print "Generating QR for {0}...".format(line),
        try:
            #set a background image with a text
            try: #set the font path for the OS
                font = ImageFont.truetype("/Library/Fonts/Arial Italic.ttf",9) # MacOS
            except:
                font = ImageFont.truetype("/usr/share/fonts/truetype/dejavu/DejaVuSerifCondensed.ttf",9) # Linux
     
            background = Image.new("RGBA", (116,126),(0,0,0))
            draw = ImageDraw.Draw(background)
            #text to be included at the bottom of the image
            draw.text((0, 117),"text prefix"+line,(255,255,255),font=font)
            draw = ImageDraw.Draw(background)
            draw = ImageDraw.Draw(background)
            #qr code layer
            img = do_qr("{0}{1}{2}".format(args.prefix, line, args.suffix),
                        mark=mark,
                        box_size=args.size,
                        border=args.border,
                        )
            background.paste(img,(0,0),img) # Overlay the QR to the bacground image

            # Save to filename.format (strip unsafe chars)
            try:
                # Sanitize filename
                filename = "".join(c for c in line if c.isalnum() or c in [' ', '-', '_', '.']).rstrip()
                filename += '.{0}'.format(args.format.strip())

                # Write out image
                background.save(filename)
            except Exception as x:
                print "!!! Could not save %s!" % filename, x
            #else:
            #   print "saved as {0}".format(filename)
        except Exception as x:
            print "ERROR while generating '%s':" % line, x
    # get the file list into a word document
    files = os.listdir(os.getcwd())
    document = Document()
    #set document margins
    sections = document.sections
    for section in sections:
        section.top_margin = Cm(0.59)
        section.bottom_margin = Cm(0.59)
        section.left_margin = Cm(1.0)
        section.right_margin = Cm(1.0)
    try:
        tbl = document.add_table(rows=1, cols=3)
        tbl.autofit = False
        tbl.style='Table Grid'
        tbl.columns[0].width = Cm(8.76)
        tbl.columns[1].width = Cm(0.71)
        tbl.columns[2].width = Cm(8.76)
    except:
        print('Error setting table format')

    
    for i in files:
        if i[len(i)-3:len(i)].upper() == 'PNG':
            row_cells = tbl.add_row().cells
            p = row_cells[0].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER            
            r = p.add_run()
            r.add_picture(i,width=Cm(3.07))

    os.chdir('../')
    shutil.rmtree(subdir)
    def remove_row(table, row):
        tbl = table._tbl
        tr = row._tr
        tbl.remove(tr)
    row = tbl.rows[0]
    remove_row(tbl, row)
    document.save('doc_name %s-%s.docx' %(start,stop))
 
    print('------ %.3f sec ------'%(time.time()-start_time))    
    print "Done!"